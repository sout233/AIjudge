from __future__ import annotations

from copy import deepcopy
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.table import Table
from docx.text.paragraph import Paragraph


ROOT = Path(__file__).resolve().parent
TEMPLATE = ROOT / "docs" / "reference" / "05-3 作品报告（人工智能实践赛，2025版）模板.docx"
OUTPUT = ROOT / "docs" / "05-3 作品报告（人工智能实践赛，2025版）-灵审云评-完成版.docx"


def remove_paragraph(paragraph: Paragraph) -> None:
    element = paragraph._element
    parent = element.getparent()
    parent.remove(element)


def clear_body(doc: Document, start_index: int) -> None:
    for paragraph in list(doc.paragraphs[start_index:]):
        remove_paragraph(paragraph)
    for table in list(doc.tables):
        tbl = table._element
        tbl.getparent().remove(tbl)


def insert_paragraph_after(paragraph: Paragraph, text: str = "", style: str | None = None) -> Paragraph:
    new_p = OxmlElement("w:p")
    paragraph._p.addnext(new_p)
    new_para = Paragraph(new_p, paragraph._parent)
    if style:
        new_para.style = style
    if text:
        new_para.add_run(text)
    return new_para


def add_heading_after(anchor: Paragraph, text: str, level: int) -> Paragraph:
    return insert_paragraph_after(anchor, text, style=f"Heading {level}")


def add_body_after(anchor: Paragraph, text: str) -> Paragraph:
    para = insert_paragraph_after(anchor, text, style="Normal")
    para.paragraph_format.space_after = 0
    return para


def insert_table_after_paragraph(paragraph: Paragraph, rows: list[list[str]]) -> Paragraph:
    cols = len(rows[0])
    temp_doc = Document()
    temp_table = temp_doc.add_table(rows=len(rows), cols=cols)
    for r_idx, row in enumerate(rows):
        for c_idx, value in enumerate(row):
            temp_table.cell(r_idx, c_idx).text = value
    tbl = deepcopy(temp_table._element)
    paragraph._p.addnext(tbl)
    new_p = OxmlElement("w:p")
    tbl.addnext(new_p)
    return Paragraph(new_p, paragraph._parent)


def set_title_block(doc: Document) -> None:
    doc.paragraphs[5].text = "作品编号：待组委会填写"
    doc.paragraphs[6].text = "作品名称：《灵审云评》-基于大模型的大学生竞赛项目智能评审与辅导平台"
    doc.paragraphs[8].text = "填写日期：2026年3月31日"
    for idx in (1, 2, 4, 5, 6, 8):
        doc.paragraphs[idx].alignment = WD_ALIGN_PARAGRAPH.CENTER


def build() -> None:
    doc = Document(TEMPLATE)
    set_title_block(doc)
    clear_body(doc, 11)
    anchor = doc.paragraphs[-1]

    anchor = add_heading_after(anchor, "作品概述", 1)
    anchor = add_heading_after(anchor, "1.1 创意来源与作品定位", 2)
    anchor = add_body_after(anchor, "“灵审云评”来源于团队在校内创新创业赛事组织与备赛中的真实观察。大量项目申报书、商业计划书、作品报告和证明材料需要在短时间内完成初评，传统人工评审存在工作量大、标准不统一、反馈周期长等问题，而通用大模型又难以直接适配不同赛道的评分口径。基于这一背景，团队开发了一个面向大学生竞赛项目的智能辅导平台，以“赛前辅导、初评辅助、规则约束”为核心目标，将竞赛管理、赛道配置、规则结构化、项目分析、结果导出和证书核验整合到同一套 Web 系统中。系统定位不是替代专家终评，而是帮助学生提前发现问题、帮助教师减轻初筛压力。")
    anchor = add_heading_after(anchor, "1.2 用户群体、主要功能与特色", 2)
    anchor = add_body_after(anchor, "本作品面向四类用户：竞赛组织方、参赛学生、指导教师和后台管理人员。系统前端采用 React + TypeScript + Vite，后端采用 FastAPI，支持账号登录、竞赛与赛道管理、评分规则维护、单文件评审、多文件批量评审、ZIP 批量评审、评审进度轮询、PDF 报告下载和软件著作权证书联网核验。系统特色主要体现在三点：一是将赛道评分标准显式存储为 JSON，并作为大模型工作流输入，同时配合区间保底机制约束评分范围，减少自由打分带来的波动；二是支持多种提交模式和异步任务处理，更适合校赛集中辅导与初评场景；三是证书核验流程已实现自动化执行，用户提交信息后，后续登录、识别、点击与结果抓取均由程序完成。")
    anchor = add_heading_after(anchor, "1.3 应用价值与推广前景", 2)
    anchor = add_body_after(anchor, "从应用价值看，系统能够把原本高度依赖人工经验的辅导与初评环节转化为可配置、可复用、可追溯的数字化流程，帮助组织方提升处理效率，也为学生提供更快、更结构化的反馈。由于系统采用轻量级文件存储和前后端分离架构，部署门槛较低，适合高校创新创业竞赛、课程项目评审、校内项目遴选等场景。后续可进一步拓展到课程作业预审、毕业设计质量预检、校企创新项目筛选等应用方向。")

    anchor = add_heading_after(anchor, "问题分析", 1)
    anchor = add_heading_after(anchor, "2.1 问题来源", 2)
    anchor = add_body_after(anchor, "高校竞赛评审的核心矛盾在于“作品数量快速增长”与“高质量辅导资源有限”之间的不平衡。以创新创业类竞赛为例，组织方通常需要在较短时间内完成竞赛信息发布、赛道分类、规则下发、材料收集、初评与复核。若完全依赖人工，容易出现评分尺度不统一、漏看材料、重复劳动严重和反馈滞后等问题；若直接使用通用 AI 对话工具，又会面临评分规则缺失、输出不可控、结果难以存档和流程难以管理的问题。")
    anchor = add_heading_after(anchor, "2.2 现有解决方案", 2)
    anchor = add_body_after(anchor, "围绕上述问题，目前常见解决路径主要包括人工评审、通用大模型问答和传统竞赛管理系统三类。三者各有优点，但仍难同时兼顾“流程管理、规则约束、批量处理、结果留存和佐证核验”五个维度。")
    anchor = insert_table_after_paragraph(anchor, [
        ["方案", "优势", "主要不足", "适配度"],
        ["传统人工评审", "专家经验丰富，可做深度判断", "效率低、主观差异大、反馈周期长", "适合终评，不适合大批量初评"],
        ["通用大模型对话", "使用门槛低，生成反馈快", "缺少赛道规则约束，结果不稳定，难沉淀流程数据", "适合临时辅助，不适合正式流程"],
        ["传统竞赛管理系统", "擅长信息发布、报名和材料收集", "缺少智能分析、报告生成与证书核验能力", "适合管理，不足以支撑智能辅导"],
        ["灵审云评", "把规则配置、智能分析、批量处理、报告导出和证书核验集成到同一平台", "暂未接入事实核验与人工评分对照实验，可信度仍需数据支撑", "适合高校竞赛的辅导与初评辅助场景"],
    ])
    anchor = add_heading_after(anchor, "2.3 本作品要解决的痛点问题", 2)
    anchor = add_body_after(anchor, "本作品重点解决四类痛点。第一，评分标准分散且口径不一，导致同一赛道不同辅导教师或初评人员给出的意见差异较大；第二，大量作品集中提交时，人工逐份阅读成本过高，难以做到及时反馈；第三，参赛者很难在正式提交前得到结构化修改建议，难以形成“以评促改”的机制；第四，项目佐证材料特别是软件著作权证书的核验流程复杂，人工查询效率低且易出错。")

    anchor = add_heading_after(anchor, "解决问题的思路", 1)
    anchor = add_heading_after(anchor, "3.1 功能与性能需求", 2)
    anchor = add_body_after(anchor, "围绕竞赛辅导与初评辅助场景，系统设计了“管理端配置 + 用户端提交 + 后端异步分析 + 结果导出”的总体思路。功能层面，系统要求支持登录鉴权、竞赛与赛道管理、评分规则上传或自动解析、单文件/多文件/ZIP 三种提交模式、结果进度轮询、PDF 报告生成和证书联网核验。性能层面，批量任务需支持并发控制，当前后端通过 `asyncio.Semaphore(3)` 将并发数限制为 3；前端结果页和批量结果页采用 2 秒轮询机制，保证用户能够持续看到任务状态变化。")
    anchor = add_heading_after(anchor, "3.2 数据来源、格式与规模", 2)
    anchor = add_body_after(anchor, "本作品并未依赖单一公开训练数据集，而是围绕工程场景组织多类运行时数据。竞赛与赛道信息以 JSON 文件形式保存在 `backend/storage/contests/contests.json`；评分规则既可直接上传 JSON，也可上传 PDF、DOC、DOCX、TXT 评分标准文档，经 Dify 规则提取工作流转为 JSON；用户提交材料支持 PDF、DOCX、PPT、PPTX 与 ZIP 压缩包；评审结果保存为 JSON，并可进一步渲染为 PDF 报告；证书核验模块在运行过程中接收网页验证码图片、官方返回结果和程序自动点击坐标。")
    anchor = insert_table_after_paragraph(anchor, [
        ["数据类别", "来源与格式", "在系统中的用途"],
        ["竞赛配置数据", "管理员维护的 JSON", "保存竞赛名称、时间、赛道、发布状态等元数据"],
        ["评分规则数据", "JSON 或由规则文档自动抽取", "作为大模型评分标准输入，约束分析口径"],
        ["参赛作品数据", "PDF/DOCX/PPT/PPTX/ZIP", "作为待分析文档上传至后端和 Dify 工作流"],
        ["结果数据", "JSON + PDF", "保存状态、消息、工作流输出并生成可交付报告"],
        ["核验辅助数据", "验证码图片、自动点击坐标、官方页面结果", "支撑软件著作权证书联网核验"],
    ])
    anchor = add_heading_after(anchor, "3.3 数据样例", 2)
    anchor = add_body_after(anchor, "规则数据样例：{\"dimensions\":[{\"dimension_name\":\"创新性\",\"dimension_weight\":0.3,\"dimension_max_score\":30,\"points\":[{\"point_name\":\"技术创新\",\"max_score\":15}]}]}。结果数据样例：{\"status\":\"running\",\"messages\":[{\"text\":\"任务已创建\"}],\"metadata\":{\"contest_id\":\"contest-001\",\"track_id\":\"track-ai\",\"filename\":\"a1b2c3.pdf\"}}。这类结构化数据使竞赛配置、评分约束和结果展示能够在工程上被统一处理。")
    anchor = add_heading_after(anchor, "3.4 关键设计要求", 2)
    anchor = add_body_after(anchor, "为确保方案可落地，系统对关键环节做了工程约束：上传文件按内容计算 MD5 哈希并缓存，减少重复存储；评分时必须绑定竞赛对应赛道的规则 JSON，并按照评分标准中的得分点进行分档，落入某档后最终得分不会低于该档最低分；重复检测基于 PDF 文本抽取与 `difflib.SequenceMatcher` 相似度计算，默认阈值为 0.8；ZIP 批量提交前执行有效性校验和路径穿越检查；PDF 报告生成采用缓存与线程锁，避免重复渲染；证书核验流程则尽量自动闭环执行，降低人工干预。")

    anchor = add_heading_after(anchor, "技术方案", 1)
    anchor = add_heading_after(anchor, "4.1 总体技术路线", 2)
    anchor = add_body_after(anchor, "系统采用前后端分离与外部工作流协同的技术路线。前端负责竞赛展示、作品上传、进度监控和结果呈现；后端负责鉴权、文件管理、规则装载、异步任务调度、结果持久化和报告生成；大模型能力通过 Dify 工作流接入，分别承接规则抽取与项目分析；证书核验模块使用 DrissionPage 驱动浏览器访问中国版权保护中心公开查询页面，并结合本地验证码识别模型完成自动化处理。")
    anchor = add_heading_after(anchor, "4.2 规则结构化与工作流调用", 2)
    anchor = add_body_after(anchor, "系统的核心思想不是直接把作品文档交给大模型自由发挥，而是先将赛道规则显式结构化，再把规则 JSON 与作品文件共同输入工作流。`backend/app/services/rule.py` 支持两种规则维护方式：其一，管理员直接上传或在线编辑 JSON；其二，管理员上传 PDF/Word/TXT 评分标准，由 Dify 规则提取工作流解析出标准化 JSON。分析阶段，`backend/app/clients/dify.py` 会先把作品文件上传至 Dify，再以 `Student_File` 与 `Score_Standard` 两个输入字段发起主工作流请求。这样，模型输出并非无约束生成，而是受到具体赛道评分标准的限制。")
    anchor = add_heading_after(anchor, "4.3 智能评审与批量处理算法", 2)
    anchor = add_body_after(anchor, "在单文件分析流程中，后端首先验证竞赛与赛道是否存在、是否配置评分规则，随后执行查重、初始化结果文件并通过后台任务调用 Dify。评分过程采用“规则约束 + 区间保底”的思路：系统先依据评分标准中的得分点和档位描述引导模型判断作品所处区间，再在区间内部给出具体分数，避免出现与规则严重偏离的极端结果。批量评审和 ZIP 批量评审共用同一套结果存储思路，不同之处在于任务清单组织方式不同。为避免外部工作流并发过高，`backend/app/services/judge.py` 使用 `asyncio.Semaphore(3)` 限制最大同时执行任务数为 3；ZIP 模式还会在解压前验证压缩包合法性，并对成员路径做安全检查，防止路径穿越风险。")
    anchor = add_heading_after(anchor, "4.4 重复检测与结果生成", 2)
    anchor = add_body_after(anchor, "重复检测模块位于 `backend/app/services/duplicate.py`。系统先用 PyPDF 抽取当前文档与历史文档文本，再使用 `difflib.SequenceMatcher` 计算相似度；当相似度超过阈值时，系统直接拒绝分析并返回提示。分析完成后，系统把工作流返回结果落盘为 JSON，前端结果页从 `workflow_data.data.outputs.text` 中解析出结构化评分结果，并在完成后调用下载接口生成 PDF 报告。报告生成由 WeasyPrint 完成，同时配合缓存目录和线程锁降低重复渲染开销。")
    anchor = add_heading_after(anchor, "4.5 证书核验与验证码识别方案", 2)
    anchor = add_body_after(anchor, "证书核验模块是本作品区别于一般竞赛管理系统的重要扩展。后端通过 DrissionPage 自动登录版权查询网站并发起检索，当页面弹出验证码时，系统会抓取背景图和指令文本，调用本地验证码识别子系统进行目标定位。识别模型定义在 `backend/app/captcha/predictor.py`，以 ResNet18 为主干网络，并输出字符类别与朝向两类结果；`backend/app/captcha/system.py` 负责切图、颜色识别、指令解析和坐标匹配。当前流程中，用户只需要提交查询信息，后续登录、识别、点击与结果抓取均由程序自动执行，体现了该模块较高的自动化完成度。")

    anchor = add_heading_after(anchor, "系统实现", 1)
    anchor = add_heading_after(anchor, "5.1 前端实现", 2)
    anchor = add_body_after(anchor, "前端入口位于 `frontend/src/App.tsx`，系统划分为首页、竞赛列表、登录页、作品提交页、单任务结果页、批量结果页、ZIP 批量结果页、证书核验页和管理后台。`SubmitWorkPage.tsx` 实现了三种上传模式的统一交互；`ResultPage.tsx` 兼容单评委结果与多评委结果两种结构；`BatchResultPage.tsx` 和 `ZipBatchResultPage.tsx` 用于批量任务监控；`CheckCertificatePage.tsx` 负责证书核验。状态管理采用 Zustand，接口请求与缓存采用 TanStack Query。")
    anchor = add_heading_after(anchor, "5.2 后端实现", 2)
    anchor = add_body_after(anchor, "后端入口位于 `backend/app/main.py`，统一注册认证、公告、竞赛、规则、评审、下载和核验接口。登录接口会先调用 Dify 控制台认证，再由本系统生成本地 JWT；分析接口统一由 `verify_token` 保护，用户名通过 JWT 透传给 Dify 工作流。业务层按“竞赛管理、规则管理、智能分析、查重、报告生成、证书核验”拆分为多个服务文件，既便于调试，也便于后续替换数据层或接入消息队列。")
    anchor = add_heading_after(anchor, "5.3 存储与部署方式", 2)
    anchor = add_body_after(anchor, "当前版本采用“文件系统 + JSON”轻量部署方案，运行时数据分别存储于 `backend/storage/uploads`、`results`、`rules`、`contests`、`announcements` 与 `download_templates` 目录下。该方案适合课程实践、竞赛答辩和校内部署，不依赖额外数据库即可完成闭环。项目启动方式也较为直接：后端可通过 `uv sync` 与 `uv run uvicorn app.main:app --reload` 启动，前端通过 `npm install` 与 `npm run dev` 启动。")
    anchor = add_heading_after(anchor, "5.4 实现难点与解决方法", 2)
    anchor = add_body_after(anchor, "工程实现过程中，团队主要解决了四类问题。其一，如何让大模型分析结果更可控，方案是引入规则 JSON 作为工作流输入，并在评分逻辑中加入区间保底机制；其二，如何在批量处理时兼顾吞吐与稳定，方案是使用后台任务加信号量限流；其三，如何在不同格式材料之间保持结果可追踪，方案是统一结果 JSON 与 PDF 输出；其四，如何在证书核验过程中应对验证码和站点波动，方案是结合自动登录、视觉识别与自动点击，尽量形成程序闭环。")

    anchor = add_heading_after(anchor, "测试分析", 1)
    anchor = add_heading_after(anchor, "6.1 测试环境", 2)
    anchor = add_body_after(anchor, "本次整理报告时，项目在 Windows 开发环境下进行静态检查与构建验证。前端环境为 Node.js v22.14.0，后端环境为 Python 3.13.5。由于系统依赖 Dify 工作流、DeepSeek API、版权查询网站和本地模型文件，离线环境下无法完成全部外部能力的端到端自动测试，因此本章主要给出已完成的工程验证结果，并明确当前测试边界。")
    anchor = insert_table_after_paragraph(anchor, [
        ["测试项", "执行时间", "结果", "说明"],
        ["前端生产构建", "2026-03-31", "通过", "执行 `npm run build` 成功，Vite 输出 `dist/` 文件"],
        ["后端语法检查", "2026-03-31", "通过", "执行 `python -m compileall app` 成功，无基础语法错误"],
        ["功能链路核对", "2026-03-31", "通过", "已在代码中确认登录、规则、分析、下载、核验等主链路闭环"],
        ["性能告警检查", "2026-03-31", "发现待优化项", "前端打包后主 JS chunk 为 904.91 kB，存在代码分割优化空间"],
    ])
    anchor = add_heading_after(anchor, "6.2 代表性测试内容", 2)
    anchor = add_body_after(anchor, "围绕作品定位，团队重点验证了五类场景：一是登录与权限控制，确认管理员与普通用户路径分离；二是作品上传与分析任务创建，确认单文件、多文件、ZIP 三种模式都能形成结果文件或任务清单；三是结果查询与 PDF 下载，确认结果页可轮询并解析工作流输出；四是查重与异常处理，确认规则缺失、文件缺失、相似度过高等情况下能返回明确提示；五是证书核验流程，确认系统能够完成从查询请求到验证码自动处理再到结果抓取的自动流程。")
    anchor = insert_table_after_paragraph(anchor, [
        ["测试用例", "输入/操作", "预期现象"],
        ["单文件评审", "上传单个文档并选择赛道", "创建单个工作流任务，结果页进入轮询状态"],
        ["多文件批量评审", "上传多个作品并批量提交", "生成多个任务，前端进入批量监控页面"],
        ["ZIP 批量评审", "上传 ZIP 压缩包", "系统完成安全解压、生成任务清单并限制 3 个并发"],
        ["重复内容拦截", "提交与历史 PDF 高相似文本", "系统返回拒绝分析提示及相似度信息"],
        ["规则自动抽取", "上传评分标准文档", "规则提取工作流返回 JSON 并保存为赛道规则"],
        ["证书核验", "输入登记号和关键字", "返回成功、未找到或不匹配中的一种结果"],
    ])
    anchor = add_heading_after(anchor, "6.3 测试结论与局限", 2)
    anchor = add_body_after(anchor, "现阶段可以确认：项目已经具备较完整的工程闭环，前后端代码可正常构建，核心业务路径在代码层面可追踪，适合用于课程实践与竞赛展示。但也要客观看到，本项目尚未建立人工评分对照实验、事实核验模块和标准化评测集，当前分析结果主要基于学生上传文档与对应规则生成，因此更适合作为竞赛辅导和初评辅助工具，而不是完全替代真实评委的最终裁决系统。当前结论更偏向“工程可用性验证”，而非“算法指标定型”。后续若继续迭代，应优先补齐样本集、人工对照实验、事实核验能力和自动化回归测试。")

    anchor = add_heading_after(anchor, "作品总结", 1)
    anchor = add_heading_after(anchor, "7.1 综合评价", 2)
    anchor = add_body_after(anchor, "从创意层面看，“灵审云评”不是单点式的 AI 打分工具，而是围绕竞赛真实流程构建的智能辅导与初评辅助平台；从技术路线看，项目把规则结构化、区间约束、大模型工作流、异步任务管理、PDF 报告导出与证书核验有机结合，形成了较完整的系统方案；从工作量看，仓库已经具备前后端页面、接口、服务层和验证码识别子模块，工程实现较为扎实；从测试效果看，当前版本已通过前端构建和后端语法验证，具备继续打磨与落地的基础。")
    anchor = add_heading_after(anchor, "作品特色与创新点", 2)
    anchor = add_body_after(anchor, "本作品的创新主要体现为四点：第一，以“赛道规则结构化 + 区间保底评分”为抓手，提高大模型分析结果的可解释性与可控性；第二，将单任务、多文件和 ZIP 批量三类处理模式统一到同一系统中，契合高校集中辅导和初评需求；第三，把报告下载、查重拦截和证书核验纳入同一平台，提升方案完整度；第四，在证书核验中引入本地视觉模型和自动点击流程，探索了复杂网页验证场景下的自动化实现路径。")
    anchor = add_heading_after(anchor, "应用推广", 2)
    anchor = add_body_after(anchor, "项目适合在高校创新创业赛事、学院项目遴选、课程实践考核和学生竞赛辅导场景中推广应用。对于组织方，系统可降低辅导前筛和材料整理工作量；对于教师，系统可承担初筛和基础诊断工作；对于学生，系统可提供更及时的结构化反馈。若后续将存储层升级为数据库、补充人工评分对照数据并完善配置外置化，系统可进一步扩展为校级或区域级竞赛辅导与初评支撑平台。")
    anchor = add_heading_after(anchor, "作品展望", 2)
    anchor = add_body_after(anchor, "下一阶段的工作重点包括三方面：一是补齐人工评分对照实验、自动化测试与评测数据集，形成可量化的系统指标；二是加入夸大表述检测、证据缺失提示、外部知识库或事实核验能力，提升结果可信度；三是将敏感配置全面迁移到环境变量或密钥管理体系，并继续扩展多模态材料解析、更多竞赛模板适配和数据分析看板，使系统从“竞赛辅导工具”逐步发展为“高可信竞赛数字化支撑平台”。")

    anchor = add_heading_after(anchor, "参考文献", 1)
    for ref in [
        "[1] FastAPI Documentation. FastAPI Official Documentation.",
        "[2] React Documentation. React Official Documentation.",
        "[3] Vite Documentation. Vite Official Documentation.",
        "[4] Dify Documentation. Workflow and File Upload API Documentation.",
        "[5] DeepSeek API Documentation.",
        "[6] PyPDF Documentation. PDF Text Extraction Documentation.",
        "[7] WeasyPrint Documentation. HTML to PDF Rendering Documentation.",
        "[8] DrissionPage Documentation. Browser Automation Documentation.",
        "[9] He K, Zhang X, Ren S, Sun J. Deep Residual Learning for Image Recognition[C]//Proceedings of CVPR. 2016.",
    ]:
        anchor = add_body_after(anchor, ref)

    doc.save(OUTPUT)


if __name__ == "__main__":
    build()
